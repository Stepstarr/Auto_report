# -*- coding: utf-8 -*-
# @Time    : 2024/3/6 12:52
# @Author  : Stepstar
# @FileName: output.py
# @Software: PyCharm
from .format_get import get_format
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
#re_format = get_format()

# content: 字典类型，key is a subtitle, value is a section.
def output(content,format,report_name,keyword=None):
    # TODO：补充自定义格式，get_format的接入
    # 创建一个Word文档对象
    document = Document()
    # 添加标题
    document.add_heading(report_name)
    for subtitle in content.keys():
        # 添加子标题
        document.add_heading(subtitle ,level=2)
        # 添加一个段落
        paragraph = document.add_paragraph(content[subtitle])
        # 通过调用get_format()设置格式
        paragraph=get_format(paragraph,keyword,format)

    document.save('example_document/example1.docx')


    #     paragraph.add_run('Lorem ipsum ')
    #     paragraph.add_run('dolor').bold = True
    #     paragraph.add_run(' sit amet.')
    #
    # # 设置段落的字体。假设我们要设置字体为宋体，字号为12号
    # run = paragraph.runs[0]
    # run.font.name = '宋体'
    # r = run._element
    # r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    #
    # # 设置字体大小
    # run.font.size = Pt(12)
    #
    # # 设置字体颜色
    # run.font.color.rgb = RGBColor(0, 0, 0)
    #
    # # 设置段落的行间距
    # paragraph_format = paragraph.paragraph_format
    # paragraph_format.line_spacing = Pt(24)  # 例如，24磅的行间距
    #
    # # 设置段落对齐方式，例如居中对齐
    # paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 保存文档
    # document.save('/mnt/data/example_document.docx')
