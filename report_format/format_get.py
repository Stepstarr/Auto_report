# -*- coding: utf-8 -*-
# @Time    : 2024/3/6 12:52
# @Author  : Stepstar
# @FileName: format_get.py
# @Software: PyCharm
from docx import Document


# 输入：paragraph：段落，keyword：关键词；输出：划分为3个runs的paragraph
def sep_paras(paragraph, keyword: str):
    text=paragraph.text
    start_index = text.find(keyword)
    new_paragraph=Document().add_paragraph()
    if start_index != -1:
        # 添加第一个run（在 keyword 之前的文本）
        first_part = text[:start_index]
        run1 = new_paragraph.add_run(first_part)
        # 添加第二个run（keyword）
        middle_part = text[start_index:start_index + len(keyword)]
        run2 = new_paragraph.add_run(middle_part)
        # 如果还有文本在 keyword 之后，添加第三个run
        if start_index + len(keyword) < len(text):
            third_part = text[start_index + len(keyword):]
            run3 = new_paragraph.add_run(third_part)
    return new_paragraph
def get_format(paragraph,keyword,format):
    if keyword != '' and keyword is not None:
        if keyword not in paragraph.text:
            raise ValueError("关键字不在段落中！")
        paragraph=sep_paras(paragraph,keyword)
    paragraph=format.set_format(paragraph)
    return paragraph
